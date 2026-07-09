"""Generate Word report for all handover DRL experiments."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None, highlight_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        set_cell_bg(cell, '2E4057')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ci == highlight_col:
                cell.paragraphs[0].runs[0].bold = True
            # Color rows alternate
            if ri % 2 == 0:
                set_cell_bg(cell, 'F0F4F8')
    return table

def add_result_table(doc, data):
    """data = list of (speed, g, pp, rlf)"""
    headers = ['Speed (km/h)', 'G_R (%)', 'PP rate (%)', 'RLF rate (%)']
    rows = []
    for speed, g, pp, rlf in data:
        rows.append([speed, f'{g:.1f}', f'{pp:.1f}', f'{rlf:.1f}'])
    add_table(doc, headers, rows, highlight_col=1)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x13, 0x78, 0x9B)

def para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE ────────────────────────────────────────────────────────────────
    title = doc.add_heading('Bao cao Tai hien Bai bao DRL Handover', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'A Deep Reinforcement Learning-Based Approach for Adaptive Handover Protocols\n'
        'SCC 2025, KIT — Ket qua toan bo thi nghiem'
    )
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Ngay bao cao: ').bold = True
    info.add_run('2026-06-24')
    info.add_run('    |    Thiet bi: ').bold = True
    info.add_run('RTX 3060 Ti + CPU')
    info.add_run('    |    Framework: ').bold = True
    info.add_run('Stable-Baselines3, PyTorch')

    doc.add_page_break()

    # ── 1. TONG QUAN ─────────────────────────────────────────────────────────
    heading(doc, '1. Tong quan bai bao va muc tieu')
    para(doc,
        'Bai bao de xuat su dung PPO (Proximal Policy Optimization) de toi uu hoa '
        'quyet dinh handover trong mang 5G voi N=5 base station. '
        'Agent quan sat SINR cua tat ca BS va chon BS de handover. '
        'Muc tieu: dat G_R (Gamma_R) ~ 99.8%, PP rate thap, RLF rate thap.'
    )

    heading(doc, '1.1 Cac thong so chinh', level=2)
    params = [
        ['t_ho_prep', '5 timesteps (500ms)', 'Thoi gian chuan bi HO'],
        ['t_ho_exec', '4 timesteps (400ms)', 'Thoi gian thuc hien HO (mat ket noi)'],
        ['Network', '[64, 128, 64] MLP', 'Kien truc PPO policy'],
        ['ent_coef', '0.1', 'He so entropy (khuyen khich explore)'],
        ['lr', '5e-5', 'Learning rate'],
        ['N steps', '5,000,000', 'Tong so buoc training'],
        ['Speeds', '30/50/70/90 km/h', 'Toc do UE trong eval'],
    ]
    add_table(doc, ['Tham so', 'Gia tri', 'Mo ta'], params)

    heading(doc, '1.2 Ket qua paper goc (muc tieu)', level=2)
    paper_data = [
        (30, 99.80, 43.58, 1.17),
        (50, 99.82, 45.38, 0.70),
        (70, 99.84, 46.45, 2.84),
        (90, 99.76, 47.53, 4.63),
    ]
    add_result_table(doc, paper_data)
    para(doc, 'Da tai hien thanh cong evaluation (sai so < 0.02%). Viec training lai tu scratch gap kho khan.',
         italic=True)

    doc.add_page_break()

    # ── 2. KET QUA BASELINE ──────────────────────────────────────────────────
    heading(doc, '2. Ket qua Baseline (tai hien thanh cong)')

    heading(doc, '2.1 3GPP Standard Baseline', level=2)
    gpp_data = [
        (30, 99.756, 41.83, 3.37),
        (50, 99.751, 50.00, 2.05),
        (70, 99.797, 47.26, 2.53),
        (90, 99.696, 44.41, 6.99),
    ]
    add_result_table(doc, gpp_data)

    doc.add_paragraph()
    heading(doc, '2.2 PPO Model goc (pre-trained)', level=2)
    ppo_data = [
        (30, 99.795, 43.58, 1.17),
        (50, 99.823, 45.38, 0.70),
        (70, 99.839, 46.45, 2.84),
        (90, 99.763, 47.53, 4.63),
    ]
    add_result_table(doc, ppo_data)
    para(doc, 'Evaluation tai hien thanh cong. Training lai gap kho khan (credit assignment problem).',
         bold=True)

    doc.add_page_break()

    # ── 3. V1 ────────────────────────────────────────────────────────────────
    heading(doc, '3. V1 — Training tu scratch (2026-06-18)')
    para(doc,
        'Phuong phap: PPO voi paper params (t_ho_prep=5, t_ho_exec=4). '
        'Ket qua: Policy stuck o uniform distribution (entropy H=1.61=ln5) xuyen suot 5M steps. '
        'Nguyen nhan: Credit assignment problem — P(HO/step)=0.16% nen reward khong phu thuoc action, '
        'gradient xap xi 0.'
    )
    v1_data = [
        (30, 22.2, 0.0, 193.3),
        (50, 10.1, 0.0, 185.7),
        (70, 16.5, 4.5, 204.5),
        (90, 29.4, 0.0, 196.2),
    ]
    add_result_table(doc, v1_data)
    para(doc, 'Ket qua: THAT BAI — Policy uniform, khong hoc duoc', bold=True)

    doc.add_page_break()

    # ── 4. V2 ────────────────────────────────────────────────────────────────
    heading(doc, '4. V2 — 4 phuong an moi (2026-06-23)')

    # 4.1
    heading(doc, '4.1 Curriculum t_ho_prep = 1 -> 3 -> 5', level=2)
    para(doc, 'Phase 0 (prep=1): entropy giam -1.61 -> -1.07 (BREAKTHROUGH). '
              'Phase 2 (prep=5): entropy ve lai -1.61 (uniform). PP=97%.')
    v2a_data = [(30,29.5,97.1,68.2),(50,12.6,95.7,64.5),(70,23.7,95.9,63.6),(90,25.3,95.7,68.9)]
    add_result_table(doc, v2a_data)

    doc.add_paragraph()
    # 4.2
    heading(doc, '4.2 Reward Shaping alpha=0.1', level=2)
    para(doc, 'Them r = 0.1*sinr_norm[action]. Entropy stuck -1.61. Alpha qua nho.')
    v2b_data = [(30,36.5,25.0,40.0),(50,22.7,41.3,173.9),(70,23.8,34.7,124.5),(90,17.8,26.3,122.8)]
    add_result_table(doc, v2b_data)

    doc.add_paragraph()
    # 4.3
    heading(doc, '4.3 Imitation Learning (BC + PPO)', level=2)
    para(doc, 'BC: 100% accuracy sau 2 epochs. PPO: policy hoc "never HO" de tranh PP penalty. '
              'PP=0%, RLF=0% nhung G thap vi UE khong theo kip BS tot nhat.')
    v2c_data = [(30,38.8,0.0,0.0),(50,41.3,0.0,0.0),(70,50.0,0.0,0.0),(90,44.3,0.0,0.0)]
    add_result_table(doc, v2c_data)

    doc.add_paragraph()
    # 4.4
    heading(doc, '4.4 Multiple Seeds (5 seeds x 2M steps)', level=2)
    para(doc, 'Tat ca 5 seeds deu uniform (H=1.609-1.610). Problem la systematic, khong phai seed-sensitive.')
    seed_rows = [['0','1.6093','Uniform'],['42','1.6094','Uniform'],['123','1.6094','Uniform'],
                 ['777','1.6093','Uniform'],['1234','1.6090','Uniform']]
    add_table(doc, ['Seed', 'Entropy cuoi', 'Trang thai'], seed_rows)

    doc.add_page_break()

    # ── 5. V3 ────────────────────────────────────────────────────────────────
    heading(doc, '5. V3 — BC + Gradual Curriculum (2026-06-24)')
    para(doc,
        'Script: train_ppo_bc_curriculum.py. '
        'BC init (30 epochs, 100% accuracy) + 5 phase curriculum '
        '(prep=1->2->3->4->5, exec=1->1->2->3->4). '
        'ent_coef=0.01 (LOI: qua thap so voi paper=0.1).'
    )

    heading(doc, '5.1 Diagnostic — BC+Curriculum Phase 0 (prep=1)', level=2)
    para(doc, 'Eval voi t_ho_prep=1, t_ho_exec=1. PP=0%, RLF=0% -> policy "never HO" ngay tu Phase 0.')
    v3a_data = [(30,38.8,0.0,0.0),(50,41.3,0.0,0.0),(70,50.0,0.0,0.0),(90,44.3,0.0,0.0)]
    add_result_table(doc, v3a_data)

    doc.add_paragraph()
    heading(doc, '5.2 BC+Curriculum Final (paper params)', level=2)
    para(doc, 'Sau 5 phase, eval voi t_ho_prep=5, t_ho_exec=4. RLF=225-295% -> HO badly.')
    v3b_data = [(30,31.3,0.0,225.0),(50,37.4,0.0,294.7),(70,25.8,0.0,235.0),(90,21.1,0.0,239.1)]
    add_result_table(doc, v3b_data)

    doc.add_page_break()

    # ── 6. V4 ────────────────────────────────────────────────────────────────
    heading(doc, '6. V4 — Cac bien the BC+Curriculum (2026-06-24)')
    para(doc,
        'Sua loi chinh: ent_coef=0.1 (paper value, thay vi 0.01). '
        'Su dung SubprocVecEnv n_envs=4 + GPU de tang toc. '
        '4 bien the voi cac thay doi khac nhau.'
    )

    heading(doc, '6.1 bc_paper — BC + paper params truc tiep (khong curriculum)', level=2)
    para(doc, 'BC init + PPO voi paper params (t_ho_prep=5, lr=5e-5, ent=0.1). '
              'entropy dao dong -1.34 -> -1.54 (co cau truc). PP=0% nhung RLF=250-490%.')
    v4a_data = [(30,24.1,0.0,250.0),(50,33.2,0.0,490.0),(70,31.1,0.0,325.0),(90,30.6,0.0,361.5)]
    add_result_table(doc, v4a_data)

    doc.add_paragraph()
    heading(doc, '6.2 bc_highent — BC + Curriculum + ent_coef=0.1', level=2)
    para(doc, 'Curriculum 5 phase voi ent_coef=0.1. G thap nhat V4. Co PP=14-24% o 50-70km/h.')
    v4b_data = [(30,9.1,0.0,225.0),(50,21.8,14.3,285.7),(70,26.6,23.8,228.6),(90,27.9,5.3,278.9)]
    add_result_table(doc, v4b_data)

    doc.add_paragraph()
    heading(doc, '6.3 bc_noterm — BC + Curriculum + khong terminate', level=2)
    para(doc, 'terminate_on_rlf=False, terminate_on_pp=False xuyen suot. '
              'Policy hoc "never HO" — giong Imitation V2. Episodes dai hon, hoc tu penalty.')
    v4c_data = [(30,38.8,0.0,0.0),(50,41.3,0.0,0.0),(70,50.0,0.0,0.0),(90,44.3,0.0,0.0)]
    add_result_table(doc, v4c_data)

    doc.add_paragraph()
    heading(doc, '6.4 bc_2m — BC + Curriculum + 2M steps/phase (10M tong)', level=2)
    para(doc, 'Tang steps 2x. Ket qua giong bc_paper — them steps khong giai quyet van de co ban.')
    v4d_data = [(30,22.9,0.0,250.0),(50,35.1,0.0,490.0),(70,31.1,0.0,350.0),(90,30.6,0.0,361.5)]
    add_result_table(doc, v4d_data)

    doc.add_page_break()

    # ── 7. TONG HOP ──────────────────────────────────────────────────────────
    heading(doc, '7. Tong hop tat ca phuong an')

    summary_headers = ['#', 'Method', 'G_R TB (%)', 'PP TB (%)', 'RLF TB (%)', 'Trang thai']
    summary_rows = [
        ['0', 'Paper PPO (original)', '99.81', '45.7', '2.3', 'TARGET'],
        ['0', '3GPP baseline', '99.75', '45.9', '3.7', 'Reference'],
        ['1', 'PPO 2-phase (V1)', '19.6', '1.1', '195.0', 'FAIL - uniform'],
        ['2', 'Curriculum 1->3->5', '22.8', '96.1', '66.3', 'FAIL - PP=97%'],
        ['3', 'Reward shaping a=0.1', '25.2', '31.8', '115.0', 'FAIL - uniform'],
        ['4', 'Imitation (BC+PPO)', '43.6', '0.0', '0.0', 'PARTIAL - never HO'],
        ['5', 'Multiseed (5 seeds)', '20.7', '49.4', '93.2', 'FAIL - uniform'],
        ['6', 'BC+Curriculum ent=0.01', '28.9', '0.0', '248.5', 'FAIL - HO badly'],
        ['7', 'bc_paper', '29.8', '0.0', '356.6', 'FAIL - HO badly'],
        ['8', 'bc_highent', '21.4', '10.9', '254.6', 'FAIL - HO badly'],
        ['9', 'bc_noterm', '43.5', '0.0', '0.0', 'PARTIAL - never HO'],
        ['10', 'bc_2m (2M/phase)', '29.9', '0.0', '362.9', 'FAIL - HO badly'],
    ]
    add_table(doc, summary_headers, summary_rows)

    doc.add_page_break()

    # ── 8. KET LUAN ──────────────────────────────────────────────────────────
    heading(doc, '8. Ket luan va Phan tich')

    heading(doc, '8.1 Hai trang thai policy pho bien', level=2)
    para(doc,
        'Sau 11 thi nghiem, tat ca policy rơi vao 1 trong 2 trang thai:'
    )
    pattern_rows = [
        ['"Never HO"', '~43%', '0%', '0%',
         'Imitation, bc_noterm',
         'An toan, khong theo kip BS tot nhat'],
        ['"HO badly"', '~25%', '~3%', '~350%',
         'bc_paper, bc_2m, bc+curriculum',
         '9-step delay -> target BS thay doi -> RLF'],
    ]
    add_table(doc,
              ['Pattern', 'G_R', 'PP', 'RLF', 'Xuat hien khi', 'Nguyen nhan'],
              pattern_rows)

    doc.add_paragraph()
    heading(doc, '8.2 Root cause chinh', level=2)
    para(doc,
        'Voi t_ho_prep=5 + t_ho_exec=4 = 9 timesteps delay (900ms), '
        'UE di chuyen 30-90 km/h co the di chuyen 7.5-22.5 met trong thoi gian do. '
        'Target BS da khong con tot nhat khi HO hoan thanh -> RLF. '
        'De thanh cong, policy can DU DOAN BS tot nhat sau 9 buoc tuong lai.'
    )

    heading(doc, '8.3 Tai sao paper model hoat dong tot?', level=2)
    para(doc,
        '(1) Co the duoc khoi tao may man voi seed dac biet.\n'
        '(2) Co the du doan duoc pattern SINR trong data cu the.\n'
        '(3) Co the co trick training khong duoc document trong paper.\n'
        'Can phan tich them behavior cua paper model de hieu mechanism.'
    )

    heading(doc, '8.4 Huong nghien cuu tiep theo', level=2)
    next_rows = [
        ['1', 'SINR look-ahead',
         'Them SINR[t+1..t+9] vao observation de policy thay truoc',
         'Cao'],
        ['2', 'Phan tich paper model',
         'So do HO cua paper model: timing, tan suat, do chinh xac',
         'Trung binh'],
        ['3', 'Future reward shaping',
         'r += alpha * sinr_norm[action, t+9] (phan thuong du bao dung)',
         'Trung binh'],
        ['4', 'Kiem tra bug env',
         'So sanh env implementation voi paper description chi tiet',
         'Cao'],
    ]
    add_table(doc, ['#', 'Huong', 'Mo ta', 'Muc do uu tien'], next_rows)

    # ── SAVE ─────────────────────────────────────────────────────────────────
    out_path = os.path.join(ROOT, 'SUMMARY_handover_DRL_experiments.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return 0


if __name__ == '__main__':
    main()
